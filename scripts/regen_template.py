#!/usr/bin/env python3
"""Rigenera templates/modulo_template.docx da templates/modulo_originale.docx.

Sostituisce SOLO i campi compilabili (trattini, date d'esempio, premio) con
segnaposto {{...}}, preservando font/formato. Da rieseguire se il modulo
originale cambia.  Requisiti: pip install python-docx

    python3 scripts/regen_template.py
"""
import re
import docx
from docx.oxml.ns import qn

SRC = 'templates/modulo_originale.docx'
DST = 'templates/modulo_template.docx'


def main():
    d = docx.Document(SRC)
    P = d.paragraphs

    def set_run(pi, ri, txt):
        P[pi].runs[ri].text = txt

    # Cognome Nome ___  targa ___
    set_run(10, 1, '{{cognome_nome}}'); set_run(10, 5, '{{targa}}')
    # Residenza ___  Città ___  CAP ___
    set_run(11, 1, '{{residenza}}'); set_run(11, 6, '{{citta}}'); set_run(11, 8, '{{cap}}')
    # Provincia __  Codice Fiscale ___  ...  tel ___
    set_run(12, 1, '{{provincia}}'); set_run(12, 3, '{{codice_fiscale}}'); set_run(12, 8, '{{tel}}')
    # email: trattini dentro l'hyperlink (mailto)
    for hl in P[12]._p.findall(qn('w:hyperlink')):
        for t in hl.iter(qn('w:t')):
            if t.text and set(t.text) <= set('_'):
                t.text = '{{email}}'
    # Date copertura
    set_run(14, 3, '{{data_inizio}}')
    for i in (4, 5, 6, 7): set_run(14, i, '')
    set_run(14, 13, '{{data_fine}}')
    for i in (14, 15, 16, 17, 18): set_run(14, i, '')
    # Premio
    set_run(17, 2, '{{premio}}')
    for i in (3, 4): set_run(17, i, '')

    d.save(DST)

    d2 = docx.Document(DST)
    full = "\n".join(p.text for p in d2.paragraphs)
    placeholders = sorted(set(re.findall(r'{{[a-z_]+}}', full)))
    leftover = [p.text for p in d2.paragraphs if re.search(r'_{3,}', p.text)]
    print('Segnaposto:', placeholders)
    print('Trattini residui:', leftover or 'nessuno')


if __name__ == '__main__':
    main()
